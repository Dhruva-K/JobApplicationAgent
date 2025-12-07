"""
Writer Agent: Generates personalized resumes and cover letters.
"""

import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime

from graph.memory import GraphMemory
from core.config import Config
from core.user_profile import UserProfile
from llm.llm_client import LLMClient
from llm.prompts import PromptTemplates
from agents.base_agent import BaseAgent

logger = logging.getLogger(__name__)


class WriterAgent(BaseAgent):
    """Agent responsible for generating personalized documents."""

    def __init__(
        self,
        graph_memory: GraphMemory,
        config: Config,
        user_profile: Optional[UserProfile] = None,
    ):
        """Initialize Writer Agent.

        Args:
            graph_memory: GraphMemory instance
            config: Configuration object
            user_profile: Optional UserProfile instance
        """
        super().__init__(name="WriterAgent", graph_memory=graph_memory, role="writer")
        self.config = config
        self.user_profile = user_profile or UserProfile(graph_memory)

        # Initialize LLM client
        llm_config = config.get_llm_config()
        self.llm_client = LLMClient(llm_config)

        logger.info(
            f"[WriterAgent] Initialized with {llm_config.get('provider')} LLM: {llm_config.get('model_name')}"
        )

    def generate_cover_letter(
        self,
        user_id: str,
        job_id: str,
        match_insights: Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        """Generate a personalized cover letter.

        Args:
            user_id: User identifier
            job_id: Job identifier
            match_insights: Optional match insights from MatcherAgent (score, strengths, concerns)

        Returns:
            Generated cover letter text or None if generation fails
        """
        # Get user profile
        user_data = self.user_profile.get_profile(user_id)
        if not user_data:
            logger.warning(f"[WriterAgent] User profile not found: {user_id}")
            return None

        # Get user resume
        resume_text = self.user_profile.get_resume(user_id) or ""

        # Get job information
        job = self.graph_memory.get_job(job_id)
        if not job:
            logger.warning(f"[WriterAgent] Job not found: {job_id}")
            return None

        # Get match insights if not provided
        if not match_insights:
            match_insights = self._get_match_insights(user_id, job_id)

        # Prepare prompt with match insights
        prompt = self._build_cover_letter_prompt(
            user_data=user_data,
            resume_text=resume_text,
            job=job,
            match_insights=match_insights,
        )

        try:
            logger.info(
                f"[WriterAgent] Generating cover letter for user {user_id} and job {job_id}"
            )
            cover_letter = self.llm_client.generate(prompt)
            logger.info(f"[WriterAgent] Successfully generated cover letter")
            return cover_letter.strip()
        except Exception as e:
            logger.error(f"[WriterAgent] Error generating cover letter: {e}")
            return None

    def generate_tailored_resume(
        self,
        user_id: str,
        job_id: str,
        match_insights: Optional[Dict[str, Any]] = None,
    ) -> Optional[str]:
        """Generate a complete tailored resume for a specific job.

        Args:
            user_id: User identifier
            job_id: Job identifier
            match_insights: Optional match insights from MatcherAgent

        Returns:
            Generated tailored resume text or None if generation fails
        """
        # Get user profile
        user_data = self.user_profile.get_profile(user_id)
        if not user_data:
            logger.warning(f"[WriterAgent] User profile not found: {user_id}")
            return None

        # Get original resume
        original_resume = self.user_profile.get_resume(user_id)
        if not original_resume:
            logger.warning(f"[WriterAgent] No resume found for user {user_id}")
            return None

        # Get job information
        job = self.graph_memory.get_job(job_id)
        if not job:
            logger.warning(f"[WriterAgent] Job not found: {job_id}")
            return None

        # Get match insights if not provided
        if not match_insights:
            match_insights = self._get_match_insights(user_id, job_id)

        # Build tailoring prompt
        prompt = self._build_resume_tailoring_prompt(
            original_resume=original_resume,
            user_data=user_data,
            job=job,
            match_insights=match_insights,
        )

        try:
            logger.info(
                f"[WriterAgent] Generating tailored resume for user {user_id} and job {job_id}"
            )
            tailored_resume = self.llm_client.generate(prompt)
            logger.info(f"[WriterAgent] Successfully generated tailored resume")
            return tailored_resume.strip()
        except Exception as e:
            logger.error(f"[WriterAgent] Error generating tailored resume: {e}")
            return None

    def export_to_text(
        self, user_id: str, job_id: str, output_dir: str = "outputs"
    ) -> Dict[str, str]:
        """Generate and export resume and cover letter as text files.

        Args:
            user_id: User identifier
            job_id: Job identifier
            output_dir: Directory to save files (default: "outputs")

        Returns:
            Dictionary with file paths: {"resume": path, "cover_letter": path}
        """
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)

        # Get job info for filename
        job = self.graph_memory.get_job(job_id)
        if not job:
            logger.error(f"[WriterAgent] Job not found: {job_id}")
            return {}

        # Get match insights
        match_insights = self._get_match_insights(user_id, job_id)

        # Create safe filename
        job_title = job.get("title", "job").replace("/", "-").replace("\\", "-")
        company_name = (
            job.get("company_name", "company").replace("/", "-").replace("\\", "-")
        )
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"{company_name}_{job_title}_{timestamp}"

        files = {}

        # Generate and save resume
        try:
            resume = self.generate_tailored_resume(user_id, job_id, match_insights)
            if resume:
                resume_path = output_path / f"{base_name}_resume.txt"
                resume_path.write_text(resume, encoding="utf-8")
                files["resume"] = str(resume_path)
                logger.info(f"[WriterAgent] Saved resume to {resume_path}")
        except Exception as e:
            logger.error(f"[WriterAgent] Failed to export resume: {e}")

        # Generate and save cover letter
        try:
            cover_letter = self.generate_cover_letter(user_id, job_id, match_insights)
            if cover_letter:
                cover_letter_path = output_path / f"{base_name}_cover_letter.txt"
                cover_letter_path.write_text(cover_letter, encoding="utf-8")
                files["cover_letter"] = str(cover_letter_path)
                logger.info(f"[WriterAgent] Saved cover letter to {cover_letter_path}")
        except Exception as e:
            logger.error(f"[WriterAgent] Failed to export cover letter: {e}")

        return files

    def export_to_docx(
        self, user_id: str, job_id: str, output_dir: str = "outputs"
    ) -> Dict[str, str]:
        """Generate and export resume and cover letter as DOCX files.

        Args:
            user_id: User identifier
            job_id: Job identifier
            output_dir: Directory to save files

        Returns:
            Dictionary with file paths: {"resume": path, "cover_letter": path}
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error(
                "[WriterAgent] python-docx not installed. Install with: pip install python-docx"
            )
            return {}

        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)

        # Get job info
        job = self.graph_memory.get_job(job_id)
        if not job:
            logger.error(f"[WriterAgent] Job not found: {job_id}")
            return {}

        # Get match insights
        match_insights = self._get_match_insights(user_id, job_id)

        # Create safe filename
        job_title = job.get("title", "job").replace("/", "-").replace("\\", "-")
        company_name = (
            job.get("company_name", "company").replace("/", "-").replace("\\", "-")
        )
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"{company_name}_{job_title}_{timestamp}"

        files = {}

        # Generate and save resume as DOCX
        try:
            resume_text = self.generate_tailored_resume(user_id, job_id, match_insights)
            if resume_text:
                doc = Document()

                # Set margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)

                # Add content
                for paragraph in resume_text.split("\n"):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.style.font.size = Pt(11)

                resume_path = output_path / f"{base_name}_resume.docx"
                doc.save(str(resume_path))
                files["resume"] = str(resume_path)
                logger.info(f"[WriterAgent] Saved resume to {resume_path}")
        except Exception as e:
            logger.error(f"[WriterAgent] Failed to export resume as DOCX: {e}")

        # Generate and save cover letter as DOCX
        try:
            cover_letter_text = self.generate_cover_letter(
                user_id, job_id, match_insights
            )
            if cover_letter_text:
                doc = Document()

                # Set margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                # Add content
                for paragraph in cover_letter_text.split("\n"):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.style.font.size = Pt(11)

                cover_letter_path = output_path / f"{base_name}_cover_letter.docx"
                doc.save(str(cover_letter_path))
                files["cover_letter"] = str(cover_letter_path)
                logger.info(f"[WriterAgent] Saved cover letter to {cover_letter_path}")
        except Exception as e:
            logger.error(f"[WriterAgent] Failed to export cover letter as DOCX: {e}")

        return files

    def _get_match_insights(self, user_id: str, job_id: str) -> Dict[str, Any]:
        """Get match insights from database.

        Args:
            user_id: User identifier
            job_id: Job identifier

        Returns:
            Match insights dictionary with score, strengths, concerns, reason
        """
        query = """
        MATCH (u:User {user_id: $user_id})-[m:MATCHED]->(j:Job {job_id: $job_id})
        RETURN m.match_score as score,
               m.match_reason as reason,
               m.strengths as strengths,
               m.concerns as concerns
        """

        results = self.graph_memory.query(query, {"user_id": user_id, "job_id": job_id})

        if results and len(results) > 0:
            result = results[0]
            return {
                "score": result.get("score", 0),
                "reason": result.get("reason", ""),
                "strengths": result.get("strengths", []),
                "concerns": result.get("concerns", []),
            }

        logger.warning(
            f"[WriterAgent] No match insights found for user {user_id} and job {job_id}"
        )
        return {"score": 0, "reason": "", "strengths": [], "concerns": []}

    def _build_cover_letter_prompt(
        self,
        user_data: Dict[str, Any],
        resume_text: str,
        job: Dict[str, Any],
        match_insights: Dict[str, Any],
    ) -> str:
        """Build cover letter generation prompt with match insights.

        Args:
            user_data: User profile data
            resume_text: User's resume text
            job: Job information
            match_insights: Match insights from MatcherAgent

        Returns:
            Formatted prompt
        """
        # Truncate resume for prompt
        resume_preview = resume_text[:2000] if resume_text else "No resume available"

        prompt = f"""You are an expert cover letter writer. Generate a compelling, personalized cover letter for the following job application.

JOB INFORMATION:
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company_name', 'N/A')}
- Description: {job.get('description', 'N/A')[:1000]}...

CANDIDATE INFORMATION:
- Name: {user_data.get('name', 'Candidate')}
- Education: {user_data.get('education_level', 'N/A')}
- Experience: {user_data.get('experience_years', 0)} years
- Resume Preview: {resume_preview}

MATCH ANALYSIS (Score: {match_insights.get('score', 0)}/100):
- Strengths: {', '.join(match_insights.get('strengths', [])[:3])}
- Areas to Address: {', '.join(match_insights.get('concerns', [])[:2])}

INSTRUCTIONS:
1. Write a professional, enthusiastic cover letter (3-4 paragraphs)
2. Highlight the candidate's STRENGTHS that align with the job
3. Address any CONCERNS by emphasizing transferable skills or willingness to learn
4. Show genuine interest in the company and role
5. Include a strong closing with call to action
6. Use professional but warm tone
7. Make it specific to THIS job, not generic

DO NOT include:
- Placeholders like [Your Name] or [Date]
- Generic statements
- Excessive flattery
- Salary expectations
- References to this prompt

Return ONLY the cover letter text, ready to use.
"""

        return prompt

    def _build_resume_tailoring_prompt(
        self,
        original_resume: str,
        user_data: Dict[str, Any],
        job: Dict[str, Any],
        match_insights: Dict[str, Any],
    ) -> str:
        """Build resume tailoring prompt with match insights.

        Args:
            original_resume: User's original resume text
            user_data: User profile data
            job: Job information
            match_insights: Match insights from MatcherAgent

        Returns:
            Formatted prompt
        """
        prompt = f"""You are an expert resume writer. Tailor the following resume for a specific job application.

ORIGINAL RESUME:
{original_resume[:2500]}

JOB INFORMATION:
- Title: {job.get('title', 'N/A')}
- Company: {job.get('company_name', 'N/A')}
- Description: {job.get('description', 'N/A')[:1000]}...
- Key Qualifications: {job.get('qualifications', 'N/A')[:500]}

MATCH ANALYSIS (Score: {match_insights.get('score', 0)}/100):
- Key Strengths: {', '.join(match_insights.get('strengths', [])[:3])}
- Gaps to Address: {', '.join(match_insights.get('concerns', [])[:2])}
- Match Reason: {match_insights.get('reason', 'N/A')[:200]}

INSTRUCTIONS:
1. Rewrite the resume to be highly tailored for THIS specific job
2. Emphasize skills and experiences that match the job requirements
3. Reorder bullet points to highlight relevant achievements first
4. Use keywords from the job description naturally
5. Quantify achievements where possible
6. Address skill gaps by:
   - Highlighting transferable skills
   - Emphasizing learning agility
   - Mentioning relevant coursework or projects
7. Keep the same truthful information - DO NOT fabricate experience
8. Maintain professional formatting with clear sections
9. Optimize for ATS (Applicant Tracking Systems)

FORMAT:
- Use clear section headers (e.g., SUMMARY, EXPERIENCE, EDUCATION, SKILLS)
- Use bullet points for achievements
- Keep it concise (aim for 1-2 pages worth of content)
- Use action verbs

Return ONLY the tailored resume text, ready to use. No explanations or meta-commentary.
"""

        return prompt

    def _save_document_to_file(
        self, document: str, user_id: str, job_id: str, document_type: str
    ) -> str:
        """Save document to file and return the file path.

        Args:
            document: Document text to save
            user_id: User identifier
            job_id: Job identifier
            document_type: Type of document ('cover_letter' or 'resume')

        Returns:
            File path where document was saved
        """
        try:
            # Get job info for filename
            job = self.graph_memory.get_job(job_id)
            if not job:
                # Fallback filename
                job_title = "job"
                company_name = "company"
            else:
                job_title = job.get("title", "job").replace("/", "-").replace("\\", "-")
                company_name = (
                    job.get("company_name", "company")
                    .replace("/", "-")
                    .replace("\\", "-")
                )

            # Create output directory
            if document_type == "cover_letter":
                output_dir = Path("outputs/cover_letters")
            else:
                output_dir = Path("outputs/resumes")

            output_dir.mkdir(parents=True, exist_ok=True)

            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_job_title = job_title[:30]  # Limit length
            safe_company = company_name[:30]
            filename = f"{safe_company}_{safe_job_title}_{timestamp}.txt"

            file_path = output_dir / filename

            # Save document
            file_path.write_text(document, encoding="utf-8")

            logger.info(f"[WriterAgent] Saved {document_type} to {file_path}")
            return str(file_path)

        except Exception as e:
            logger.error(f"[WriterAgent] Error saving document: {e}")
            # Return a fallback path
            return f"outputs/{document_type}_{job_id}.txt"

    async def _handle_data_request(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Handle document generation request from another agent.

        Args:
            payload: Request with 'user_id', 'job_id', and 'document_type'

        Returns:
            Response with generated document
        """
        try:
            user_id = payload.get("user_id")
            job_id = payload.get("job_id")
            document_type = payload.get("document_type", "cover_letter")  # or 'resume'
            match_insights = payload.get("match_insights")

            if not user_id or not job_id:
                return {"status": "error", "error": "user_id and job_id required"}

            logger.info(f"[WriterAgent] Generating {document_type} for job {job_id}")

            if document_type == "cover_letter":
                document = self.generate_cover_letter(user_id, job_id, match_insights)
            elif document_type == "resume":
                document = self.generate_tailored_resume(
                    user_id, job_id, match_insights
                )
            else:
                return {
                    "status": "error",
                    "error": f"Unknown document type: {document_type}",
                }

            if not document:
                return {
                    "status": "error",
                    "error": f"Failed to generate {document_type}",
                }

            # Save document to file
            file_path = self._save_document_to_file(
                document=document,
                user_id=user_id,
                job_id=job_id,
                document_type=document_type,
            )

            return {
                "status": "success",
                "document_type": document_type,
                "document": document,
                "file_path": file_path,
            }
        except Exception as e:
            logger.error(f"[WriterAgent] Error handling request: {e}")
            return {"status": "error", "error": str(e)}

    async def _handle_status_update(self, payload: Dict[str, Any]) -> Dict[str, Any]:
        """Handle status update from orchestrator.

        Args:
            payload: Status update details

        Returns:
            Current agent status
        """
        return {
            "status": "acknowledged",
            "agent": "writer",
            "llm_provider": self.llm_client.provider,
        }
